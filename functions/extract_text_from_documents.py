#Importing libraries

from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain_community.document_loaders import PyPDFLoader
import chardet
import xml.etree.ElementTree as ET
from docx import Document
import tempfile
import os
import shutil

#Function to extract and upload into folder
UPLOAD_FOLDER = 'uploaded_files'

def get_documents_from_files(docs):
    # Ensure the upload directory exists
    if os.path.exists(UPLOAD_FOLDER):
        shutil.rmtree(UPLOAD_FOLDER)
    os.makedirs(UPLOAD_FOLDER)

    documents = []
    for doc in docs:
        # Create a unique file path within the upload directory
        file_path = os.path.join(UPLOAD_FOLDER, doc.filename)

        # Save the uploaded file to the specified directory
        with open(file_path, 'wb') as f:
            f.write(doc.read())

        # Process the file based on its type
        if doc.filename.endswith('.pdf'):
            loader = PyPDFLoader(file_path)
            documents.extend(loader.load_and_split())
        elif doc.filename.endswith('.txt'):
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            encoding = chardet.detect(raw_data)['encoding']
            loader = TextLoader(file_path, encoding=encoding)
            documents.extend(loader.load())
        elif doc.filename.endswith('.xml'):
            tree = ET.parse(file_path)
            root = tree.getroot()
            for elem in root.iter():
                if elem.text:
                    documents.append({"page_content": elem.text + "\n"})
        elif doc.filename.endswith('.docx'):
            docx_document = Document(file_path)
            for para in docx_document.paragraphs:
                documents.append({"page_content": para.text + "\n"})
        else:
            continue

    return documents


# Function to extract text from PDFs or text documents using LangChain document loaders
# def get_documents_from_files(docs):
#     documents = []
#     for doc in docs:
#         with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#             tmp_file.write(doc.read())
#             tmp_file.flush()
#             tmp_file.seek(0)

#             if doc.filename.endswith('.pdf'):
#                 loader = PyPDFLoader(tmp_file.name)
#                 documents.extend(loader.load_and_split())
#             elif doc.filename.endswith('.txt'):
#                 with open(tmp_file.name, 'rb') as f:
#                     raw_data = f.read()
#                 encoding = chardet.detect(raw_data)['encoding']
#                 loader = TextLoader(tmp_file.name, encoding=encoding)
#                 documents.extend(loader.load())
#             elif doc.filename.endswith('.xml'):
#                 tree = ET.parse(tmp_file.name)
#                 root = tree.getroot()
#                 for elem in root.iter():
#                     if elem.text:
#                         documents.append({"page_content": elem.text + "\n"})
#             elif doc.filename.endswith('.docx'):
#                 doc = Document(tmp_file.name)
#                 for para in doc.paragraphs:
#                     documents.append({"page_content": para.text + "\n"})
#             else:
#                 continue
#     return documents


# # Function to extract text from PDFs or text documents using LangChain document loaders

# def get_text_from_documents(docs):
#     text = ""
#     for doc in docs:
#         with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#             tmp_file.write(doc.read())
#             tmp_file.flush()
#             tmp_file.seek(0)

#             if doc.filename.endswith('.pdf'):
#                 loader = PyPDFLoader(tmp_file.name)
#                 documents = loader.load_and_split()
#                 for document in documents:
#                     text += document.page_content
#             elif doc.filename.endswith('.txt'):
#                 with open(tmp_file.name, 'rb') as f:
#                     raw_data = f.read()
#                 encoding = chardet.detect(raw_data)['encoding']
#                 loader = TextLoader(tmp_file.name, encoding=encoding)
#                 documents = loader.load()
#                 for document in documents:
#                     text += document.page_content
#             elif doc.filename.endswith('.xml'):
#                 tree = ET.parse(tmp_file.name)
#                 root = tree.getroot()
#                 for elem in root.iter():
#                     if elem.text:
#                         text += elem.text + "\n"
#             elif doc.filename.endswith('.docx'):
#                 doc = Document(tmp_file.name)
#                 for para in doc.paragraphs:
#                     text += para.text + "\n"
#             else:
#                 continue
#     return text