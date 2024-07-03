from flask import Flask, request, jsonify
from flask_cors import CORS
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.memory import ConversationBufferMemory
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from docx import Document as DocxDocument
import xml.etree.ElementTree as ET
import os
import warnings
import tempfile
import chardet
from deduplication import dataDeduplication  # Import your deduplication function

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

warnings.filterwarnings("ignore", message="`resume_download` is deprecated and will be removed in version 1.0.0.")

from dotenv import load_dotenv
load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")

llm = ChatGroq(groq_api_key=groq_api_key, model='Llama3-70b-8192', temperature=0.1, max_tokens=2000)

session_state = {
    'conversation': None,
    'chat_history': [],
    'vectorstore': None,
    'raw_text': None,
    'document_references': [],
    'chunk_references': [],
    'deduplicated_text': None,
    'duplicate_text': None,
    'original_word_count': None,
    'duplicate_word_count': None,
    'chronological_response': None,
    'summarization_response': None,
}

def get_text_from_documents(docs):
    text = ""
    document_references = []

    for doc in docs:
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(doc.read())
            tmp_file.flush()
            tmp_file.seek(0)

            doc_name = doc.filename  # Get the name here

            if doc_name.endswith('.pdf'):
                loader = PyPDFLoader(tmp_file.name)
                documents = loader.load()
                for document in documents:
                    page_text = document.page_content.replace('\n', ' ')
                    text += page_text
                    document_references.append((doc_name, page_text))
            elif doc_name.endswith('.txt'):
                with open(tmp_file.name, 'rb') as f:
                    raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding']
                loader = TextLoader(tmp_file.name, encoding=encoding)
                documents = loader.load()
                for document in documents:
                    page_text = document.page_content.replace('\n', ' ')
                    text += page_text
                    document_references.append((doc_name, page_text))
            elif doc_name.endswith('.xml'):
                tree = ET.parse(tmp_file.name)
                root = tree.getroot()
                for elem in root.iter():
                    if elem.text:
                        page_text = elem.text.replace('\n', ' ')
                        text += page_text + "\n"
                        document_references.append((doc_name, page_text))
            elif doc_name.endswith('.docx'):
                docx_doc = DocxDocument(tmp_file.name)
                for para in docx_doc.paragraphs:
                    page_text = para.text.replace('\n', ' ')
                    text += page_text + "\n"
                    document_references.append((doc_name, page_text))
            else:
                continue

    return text, document_references

def get_text_chunks(text, document_references):
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=2002,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    chunk_references = []
    start_idx = 0

    for chunk in chunks:
        end_idx = start_idx + len(chunk)
        for ref in document_references:
            if ref[1] in text[start_idx:end_idx]:
                chunk_references.append({'source': ref[0], 'content': chunk})
                break
        start_idx = end_idx

    return chunk_references

class Document:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata

def get_vectorstore(chunk_references):
    embeddings = HuggingFaceEmbeddings()
    docs_with_metadata = [Document(chunk['content'], {'source': chunk['source']}) for chunk in chunk_references]
    vectorstore = FAISS.from_documents(docs_with_metadata, embedding=embeddings)
    return vectorstore

def get_conversation_chain(vectorstore):
    memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
    conversation_chain = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=vectorstore.as_retriever(),
        memory=memory,
        output_key='result',
    )
    return conversation_chain

@app.route('/upload', methods=['POST'])
def handle_upload():
    docs = request.files.getlist("files")
    if not docs:
        return jsonify({"error": "No files uploaded."}), 400

    raw_text, document_references = get_text_from_documents(docs)
    session_state['raw_text'] = raw_text
    session_state['document_references'] = document_references

    chunk_references = get_text_chunks(raw_text, document_references)
    session_state['chunk_references'] = chunk_references

    vectorstore = get_vectorstore(chunk_references)
    session_state['vectorstore'] = vectorstore

    session_state['conversation'] = get_conversation_chain(vectorstore)

    return jsonify({"message": "Documents processed successfully!"}), 200

@app.route('/ask', methods=['POST'])
def handle_userinput():
    data = request.json
    user_question = data.get("question")

    if not user_question:
        return jsonify({"error": "No question provided."}), 400

    if session_state['conversation'] is None:
        return jsonify({"error": "Please upload and process the documents first."}), 400

    retriever = session_state['vectorstore'].as_retriever()
    relevant_docs = retriever.get_relevant_documents(user_question)

    if not relevant_docs:
        return jsonify({"error": "No relevant documents found."}), 400

    context = " ".join([doc.page_content for doc in relevant_docs])
    prompt = generate_prompt(user_question, context)

    response = session_state['conversation']({'query': prompt})

    if 'result' in response:
        response_with_references, references = add_references_to_response(response['result'])

        session_state['chat_history'].append({"role": "user", "content": user_question})
        session_state['chat_history'].append({"role": "assistant", "content": response_with_references, "references": references})

        return jsonify({"response": response_with_references}), 200
    else:
        return jsonify({"error": "Unable to retrieve context from the response."}), 500

def add_references_to_response(response_text):
    references = []
    response_with_references = response_text

    for ref in session_state['chunk_references']:
        if ref['content'] in response_text:
            references.append(ref['source'])
    
    if references:
        response_with_references += "\n\nReferences:\n" + "\n".join(references)

    return response_with_references, references

def generate_prompt(question, context):
    return f"""
    Answer the questions based on the provided context only.
    Please provide the most accurate response based on the question.
    <context>
    {context}
    </context>
    Questions: {question}
    """

@app.route('/deduplication', methods=['POST'])
def handle_deduplication():
    if "raw_text" not in session_state:
        return jsonify({"error": "Please upload and process the documents first."}), 400
    if session_state['conversation'] is None:
        return jsonify({"error": "Please upload and process the documents first."}), 400

    deduplicated_text, duplicate_text, original_word_count, duplicate_word_count = dataDeduplication(session_state['raw_text'])
    session_state['deduplicated_text'] = deduplicated_text
    session_state['duplicate_text'] = duplicate_text
    session_state['original_word_count'] = original_word_count
    session_state['duplicate_word_count'] = duplicate_word_count

    return jsonify({"deduplicated_text": deduplicated_text, "duplicate_text": duplicate_text}), 200
    
@app.route('/chronological', methods=['POST'])
def handle_chronological_order():
    data = request.json
    query = data.get("query")

    if not query:
        return jsonify({"error": "No query provided"}), 400

    if session_state['conversation'] is None:
        return jsonify({"error": "Please upload and process the documents first."}), 400

    response = session_state['conversation']({'query': query})
    session_state['chat_history'].append({"role": "user", "content": query})
    session_state['chat_history'].append({"role": "assistant", "content": response['result']})
    session_state['chronological_response'] = response['result']

    return jsonify({"response": response['result']}), 200

@app.route('/summarization', methods=['POST'])
def handle_summarization():
    data = request.json
    query = data.get("query")

    if not query:
        return jsonify({"error": "No query provided"}), 400

    if session_state['conversation'] is None:
        return jsonify({"error": "Please upload and process the documents first."}), 400

    response = session_state['conversation']({'query': query})
    session_state['chat_history'].append({"role": "user", "content": query})
    session_state['chat_history'].append({"role": "assistant", "content": response['result']})
    session_state['summarization_response'] = response['result']

    return jsonify({"response": response['result']}), 200

if __name__ == '__main__':
    app.run(debug=True)
